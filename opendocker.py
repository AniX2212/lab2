<<<<<<< HEAD
import docx
import os

# Визначаємо шлях до файлу
file_path = os.path.join(os.getcwd(), 'Docker Basics.docx')

# Відкриваємо файл
document = docx.Document(file_path)

# Виводимо текст документу в термінал
for paragraph in document.paragraphs:
    print(paragraph.text)
=======
import docx
import os

# Визначаємо шлях до файлу
file_path = os.path.join(os.getcwd(), 'Docker Basics.docx')

# Відкриваємо файл
document = docx.Document(file_path)

# Виводимо текст документу в термінал
for paragraph in document.paragraphs:
    print(paragraph.text)
>>>>>>> 8dcfd5a5c15720480a25c4996abe6461aea1b6c0
